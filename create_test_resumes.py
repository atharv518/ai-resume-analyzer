import os
import docx

os.makedirs("test_files", exist_ok=True)

# 1. Fresher Resume
doc1 = docx.Document()
doc1.add_heading("Alex Mercer", level=1)
doc1.add_paragraph("alex.mercer@example.com | +1 (555) 234-5678 | San Francisco, CA")
doc1.add_heading("Education", level=2)
doc1.add_paragraph("B.S. in Computer Science - University of California, Berkeley (2020 - 2024)")
doc1.add_heading("Technical Skills", level=2)
doc1.add_paragraph("Languages & Frameworks: Python, FastAPI, JavaScript, React, SQL, HTML, CSS\nTools & Platforms: Docker, PostgreSQL, Git, GitHub, REST APIs, Linux")
doc1.add_heading("Projects", level=2)
doc1.add_paragraph("AI ATS Resume Analyzer (2024)\nBuilt a full-stack automated document analysis platform using FastAPI, Python, and React. Implemented PDF parsing with pypdf and achieved 95% keyword extraction accuracy.")
doc1.add_paragraph("Microservices E-Commerce API (2023)\nDesigned and deployed high-throughput RESTful backend services using FastAPI and PostgreSQL, handling 250+ req/sec with sub-40ms latency.")
doc1.add_heading("Certifications", level=2)
doc1.add_paragraph("AWS Certified Cloud Practitioner (2024)")
doc1.save("test_files/fresher_resume.docx")

# 2. Experienced Resume
doc2 = docx.Document()
doc2.add_heading("Sarah Jenkins", level=1)
doc2.add_paragraph("sarah.jenkins@example.com | +1 (555) 987-6543 | New York, NY")
doc2.add_heading("Professional Experience", level=2)
doc2.add_paragraph("Senior Software Engineer - Acme Cloud Solutions (2021 - 2024)\nArchitected and deployed enterprise microservices in Python, FastAPI, Docker, and AWS ECS. Mentored 4 junior engineers and optimized database queries in PostgreSQL, reducing query latency by 40%.")
doc2.add_paragraph("Software Developer - TechCore Systems (2018 - 2021)\nDeveloped and maintained scalable RESTful APIs with Python, Django, and PostgreSQL. Integrated Stripe payment pipelines and configured automated CI/CD workflows.")
doc2.add_heading("Education", level=2)
doc2.add_paragraph("B.S. in Computer Science - Columbia University (2014 - 2018)")
doc2.add_heading("Technical Skills", level=2)
doc2.add_paragraph("Python, FastAPI, Django, PostgreSQL, Docker, AWS, Microservices, CI/CD, Redis, Git")
doc2.save("test_files/experienced_resume.docx")

print("Created test resume files in test_files/")
