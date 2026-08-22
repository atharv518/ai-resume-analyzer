import io
from fastapi import HTTPException
from fastapi.testclient import TestClient
import docx

from app.main import app
from app.services.resume_validator import calculate_resume_scores, validate_resume_content

client = TestClient(app)


def test_resume_signals_fresher():
    fresher_text = """
    Alex Mercer
    alex.mercer@example.com | +1 (555) 234-5678 | San Francisco, CA | linkedin.com/in/alexmercer
    
    Education
    B.S. in Computer Science - University of California, Berkeley (2020 - 2024)
    
    Technical Skills
    Languages & Frameworks: Python, FastAPI, JavaScript, React, SQL, HTML, CSS
    Tools: Docker, PostgreSQL, Git, GitHub, REST APIs, Linux
    
    Projects
    AI ATS Resume Analyzer (2024)
    Built a full-stack automated document analysis platform using FastAPI, Python, and React.
    Implemented PDF parsing with pypdf and achieved 95% keyword extraction accuracy.
    
    Certifications
    AWS Certified Cloud Practitioner (2024)
    """
    pos, neg, details = calculate_resume_scores(fresher_text)
    assert pos >= 15, f"Expected pos >= 15, got {pos}"
    assert neg == 0, f"Expected neg == 0, got {neg}"
    assert details["has_candidate_name"] is True
    assert details["has_email"] is True
    assert details["has_phone"] is True
    assert len(details["detected_sections"]) >= 3

    # Validation should succeed without throwing
    validate_resume_content(fresher_text)
    print("[PASS] Fresher resume content validation passed (Score:", pos, "Neg:", neg, ").")


def test_resume_signals_experienced():
    exp_text = """
    Sarah Jenkins
    sarah.jenkins@example.com | +1 (555) 987-6543 | New York, NY | github.com/sarahjenkins
    
    Professional Experience
    Senior Software Engineer - Acme Cloud Solutions (2021 - 2024)
    Architected and deployed enterprise microservices in Python, FastAPI, Docker, and AWS ECS.
    Mentored 4 junior engineers and optimized database queries in PostgreSQL, reducing query latency by 40%.
    
    Software Developer - TechCore Systems (2018 - 2021)
    Developed and maintained scalable RESTful APIs with Python, Django, and PostgreSQL.
    
    Education
    B.S. in Computer Science - Columbia University (2014 - 2018)
    
    Technical Skills
    Python, FastAPI, Django, PostgreSQL, Docker, AWS, Microservices, CI/CD, Redis, Git
    """
    pos, neg, details = calculate_resume_scores(exp_text)
    assert pos >= 15, f"Expected pos >= 15, got {pos}"
    assert neg == 0, f"Expected neg == 0, got {neg}"
    assert details["has_candidate_name"] is True
    assert details["has_email"] is True

    # Validation should succeed without throwing
    validate_resume_content(exp_text)
    print("[PASS] Experienced resume content validation passed (Score:", pos, "Neg:", neg, ").")


def test_data_science_lab_experiment_rejection():
    # Synthetic lab experiment document structure
    dslab_text = """
    PACIFIC METROPOLITAN INSTITUTE OF TECHNOLOGY
    (Affiliated to State Technical Board, Approved by Education Council)
    100 Innovation Parkway, Suite 400, Metro City, ST-98101
    Department of Computer Science & Information Systems
    
    Data Science Lab Experiment Writing Instructions
    
    Experiment No: 03
    
    Aim: Implement Fuzzy Membership Function using python API
    
    Theory: [ Based slide shared on LMS explain following points within 2-3 pages]
    1. What is fuzzy set ?
    2. What do you mean by Membership functions w.r.t. fuzzy sets ?
    3. What are different membership function and their respective equations ?
    
    Performance: [Go to git hub repository: https://github.com/example-edu/membership-functions-lab ]
    Understand the implementation of fuzzy membership function using graphics.py and 
    membershipfunction.py by executing graphics.py write the output of the execution. Change the 
    parameters of various function and check its effect in plots.
    
    Conclusion: Thus we have learned implementation of Fuzzy Membership Function using python API.
    """
    pos, neg, details = calculate_resume_scores(dslab_text)
    assert neg >= 10, f"Expected high negative score for lab manual, got {neg}"
    
    try:
        validate_resume_content(dslab_text)
        assert False, "DSLAB_EXP_3 lab experiment should have been rejected"
    except HTTPException as exc:
        assert "The uploaded document does not appear to be a resume" in exc.detail
        print("[PASS] DSLAB_EXP_3 Lab Experiment rejection confirmed (Neg Score:", neg, "Pos Score:", pos, "):", exc.detail)


def test_resume_content_validation_rejections():
    # 1. Invoice Document
    invoice_text = """
    INVOICE
    Invoice Number: INV-2024-0091
    Date: August 12, 2024
    Bill To: Global Tech Enterprise LLC
    Payment Terms: Net 30 Days
    
    Description                  Quantity     Rate       Amount
    Web Development Consulting         40     $120.00    $4,800.00
    Server Maintenance & DevOps        10     $150.00    $1,500.00
    
    Subtotal: $6,300.00
    Tax (8%): $504.00
    Amount Due: $6,804.00
    
    Remittance Advice: Please send payments via wire transfer to Acme Corp.
    """
    try:
        validate_resume_content(invoice_text)
        assert False, "Invoice should have failed resume content validation"
    except HTTPException as exc:
        assert "The uploaded document does not appear to be a resume" in exc.detail
        print("[PASS] Invoice rejection confirmed:", exc.detail)

    # 2. Homework / Question Paper Assignment
    homework_text = """
    History of Computing - Assignment 2
    Question Paper
    Student Name: John
    Date: September 2024
    Course Code: CS-101
    
    Question 1: Describe the architectural differences between RISC and CISC instruction set architectures.
    Answer the following questions within 2 pages:
    1. Explain following points regarding pipelining.
    2. What is fuzzy set logic?
    
    Conclusion:
    Modern processors often blend both approaches by decoding CISC instructions into internal RISC-like micro-operations.
    """
    try:
        validate_resume_content(homework_text)
        assert False, "Homework assignment should have failed resume content validation"
    except HTTPException as exc:
        assert "The uploaded document does not appear to be a resume" in exc.detail
        print("[PASS] Homework assignment rejection confirmed:", exc.detail)

    # 3. Short random text
    short_text = "This is a random short text note about meeting on Friday."
    try:
        validate_resume_content(short_text)
        assert False, "Short text should have failed resume content validation"
    except HTTPException as exc:
        assert "The uploaded document does not appear to be a resume" in exc.detail
        print("[PASS] Short text rejection confirmed:", exc.detail)


def test_api_general_audit_vs_jd_audit():
    # Create docx resume bytes in memory
    doc = docx.Document()
    doc.add_heading("Alex Mercer", level=1)
    doc.add_paragraph("alex.mercer@example.com | +1 (555) 234-5678 | San Francisco, CA")
    doc.add_heading("Education", level=2)
    doc.add_paragraph("B.S. in Computer Science - University of California, Berkeley (2020 - 2024)")
    doc.add_heading("Technical Skills", level=2)
    doc.add_paragraph("Python, FastAPI, JavaScript, React, SQL, HTML, CSS, Git, PostgreSQL")
    doc.add_heading("Projects", level=2)
    doc.add_paragraph("AI Resume Analyzer\nBuilt a full-stack automated document analysis platform using FastAPI, Python, and React.")

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    resume_bytes = bio.getvalue()

    # 1. Test General Profile Audit (No JD provided)
    res_general = client.post(
        "/api/analyze",
        files={"resume": ("resume.docx", resume_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        data={"job_description": ""}
    )
    assert res_general.status_code == 200, res_general.text
    data_general = res_general.json()
    assert data_general["success"] is True
    assert data_general["job_description_provided"] is False
    assert len(data_general["skill_comparison"]["matching_skills"]) > 0
    # Must NOT have missing skills or missing keywords in general audit
    assert data_general["skill_comparison"]["missing_skills"] == []
    assert data_general["skill_comparison"]["missing_keywords"] == []
    print("[PASS] General Audit API returns empty missing_skills & missing_keywords.")

    # 2. Test Target JD Audit (JD provided with missing requirements)
    jd_text = """
    Senior Cloud Engineer
    Requirements:
    - 3+ years experience with Python, FastAPI, Docker, Kubernetes, AWS, Terraform, Kafka
    - Bachelor's degree in Computer Science
    """
    res_jd = client.post(
        "/api/analyze",
        files={"resume": ("resume.docx", resume_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        data={"job_description": jd_text}
    )
    assert res_jd.status_code == 200, res_jd.text
    data_jd = res_jd.json()
    assert data_jd["success"] is True
    assert data_jd["job_description_provided"] is True
    assert len(data_jd["skill_comparison"]["matching_skills"]) > 0
    # Must have missing skills detected against JD
    assert len(data_jd["skill_comparison"]["missing_skills"]) > 0
    assert any("docker" in s.lower() or "aws" in s.lower() or "kubernetes" in s.lower() for s in data_jd["skill_comparison"]["missing_skills"])
    print("[PASS] Target JD Audit API correctly detects missing skills:", data_jd["skill_comparison"]["missing_skills"])

    # 3. Test Non-Resume Upload via API (Invoice docx)
    doc_inv = docx.Document()
    doc_inv.add_heading("INVOICE #9821", level=1)
    doc_inv.add_paragraph("Bill To: Acme Corp | Payment Terms: Net 30 Days")
    doc_inv.add_paragraph("Amount Due: $5,000.00 | Subtotal: $4,500.00 | Tax: $500.00")
    bio_inv = io.BytesIO()
    doc_inv.save(bio_inv)
    bio_inv.seek(0)

    res_inv = client.post(
        "/api/analyze",
        files={"resume": ("invoice.docx", bio_inv.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        data={"job_description": ""}
    )
    assert res_inv.status_code == 400
    assert "The uploaded document does not appear to be a resume. Please upload a valid resume." in res_inv.json()["detail"]
    print("[PASS] API non-resume rejection returned HTTP 400 with expected error detail.")

    # 4. Test Lab Experiment Upload via API (DSLAB_EXP_3 docx)
    doc_lab = docx.Document()
    doc_lab.add_heading("PACIFIC METROPOLITAN INSTITUTE OF TECHNOLOGY", level=1)
    doc_lab.add_paragraph("Department of Computer Science & Information Systems\nData Science Lab Experiment Writing Instructions")
    doc_lab.add_heading("Experiment No: 03", level=2)
    doc_lab.add_paragraph("Aim: Implement Fuzzy Membership Function using python API")
    doc_lab.add_paragraph("Theory: [ Based slide shared on LMS explain following points ]\n1. What is fuzzy set ?")
    doc_lab.add_paragraph("Performance: [Go to git hub repository: https://github.com/example-edu/membership-functions-lab ]")
    doc_lab.add_paragraph("Conclusion: Thus we have learned implementation of Fuzzy Membership Function.")
    bio_lab = io.BytesIO()
    doc_lab.save(bio_lab)
    bio_lab.seek(0)

    res_lab = client.post(
        "/api/analyze",
        files={"resume": ("DSLAB_EXP_3.docx", bio_lab.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        data={"job_description": ""}
    )
    assert res_lab.status_code == 400
    assert "The uploaded document does not appear to be a resume. Please upload a valid resume." in res_lab.json()["detail"]
    print("[PASS] API lab experiment rejection returned HTTP 400 with expected error detail.")


if __name__ == "__main__":
    print("Running Strengthened Improvements Test Suite...")
    test_resume_signals_fresher()
    test_resume_signals_experienced()
    test_data_science_lab_experiment_rejection()
    test_resume_content_validation_rejections()
    test_api_general_audit_vs_jd_audit()
    print("\nALL STRENGTHENED IMPROVEMENT TESTS PASSED SUCCESSFULLY!")
