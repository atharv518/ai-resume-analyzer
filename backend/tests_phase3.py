import io
import asyncio
from fastapi import UploadFile
from starlette.datastructures import Headers
from app.routes.analyze import analyze_resume
from app.config import get_feature_flags
from app.services.experience_detector import classify_experience_text
from app.services.job_matcher import compare_resume_with_jd
from app.services.ats_scorer import calculate_ats_score


FRESHER_RESUME_TEXT = """
Alex Mercer
alex.mercer@example.com
(555) 234-5678

EDUCATION
Bachelor of Science in Computer Science, Tech State University (2020 - 2024)

SKILLS
Python, FastAPI, JavaScript, React, PostgreSQL, Docker, Git

PROJECTS
AI Resume Analyzer
Built full-stack ATS resume analysis platform with FastAPI, React, and PostgreSQL. Implemented text extraction from PDF documents and automated keyword scoring.

E-Commerce API Service
Developed high-throughput RESTful API with FastAPI and Redis caching, serving 200+ product endpoints with under 50ms latency.

CERTIFICATIONS
AWS Certified Cloud Practitioner
"""

EXPERIENCED_RESUME_TEXT = """
Sarah Jenkins
sarah.jenkins@example.com
+1 (555) 987-6543

PROFESSIONAL EXPERIENCE
Senior Software Engineer - Acme Cloud Solutions (2021 - 2024)
Designed and deployed scalable microservices using Python, FastAPI, Docker, and AWS ECS. Optimized PostgreSQL database queries reducing response times by 35%.

Software Developer - Core Systems Inc (2018 - 2021)
Developed RESTful APIs and integrated payment gateways using Python and Django. Managed CI/CD pipelines in GitLab.

EDUCATION
B.S. in Computer Science (2014 - 2018)

SKILLS
Python, FastAPI, Docker, AWS, PostgreSQL, Microservices, CI/CD, Redis
"""

VIRTUAL_SIMULATION_RESUME_TEXT = """
Jordan Lee
jordan.lee@example.com
(555) 345-6789

EDUCATION
B.Tech in Computer Engineering (2021 - 2025)

EXPERIENCE
JPMorgan Chase & Co. Software Engineering Virtual Experience on Forage (2023)
Completed simulated tasks involving financial data feed processing, perspective open-source library integration, and TypeScript data visualization.

PROJECTS
Task Manager App - Built with React and FastAPI

SKILLS
Python, React, TypeScript, Git
"""

SAMPLE_JD = """
We are looking for a Python Developer with strong hands-on experience in FastAPI, Docker, PostgreSQL, REST APIs, and AWS.
The candidate will design scalable backend services, optimize database queries, and collaborate in an Agile environment.
"""


def test_fresher_analysis():
    exp = classify_experience_text([], ["AI Resume Analyzer", "E-Commerce API Service"], FRESHER_RESUME_TEXT)
    assert exp["candidate_type"] == "fresher", f"Expected fresher, got {exp['candidate_type']}"
    assert exp["has_professional_experience"] is False
    assert exp["include_experience_section"] is False

    match = compare_resume_with_jd(FRESHER_RESUME_TEXT, ["Python", "FastAPI", "React", "PostgreSQL", "Docker", "Git"], SAMPLE_JD)
    assert "Python" in match["matching_skills"]
    assert "FastAPI" in match["matching_skills"]
    assert "PostgreSQL" in match["matching_skills"]

    score = calculate_ats_score(
        name="Alex Mercer",
        email="alex.mercer@example.com",
        phone="(555) 234-5678",
        skills=["Python", "FastAPI", "React", "PostgreSQL", "Docker", "Git"],
        education=["Bachelor of Science in Computer Science"],
        projects=["AI Resume Analyzer", "E-Commerce API Service"],
        experience=[],
        certifications=["AWS Certified Cloud Practitioner"],
        raw_text=FRESHER_RESUME_TEXT,
        jd_text=SAMPLE_JD,
        match_results=match,
        exp_classification=exp,
    )
    assert 0 <= score["overall_score"] <= 100
    assert score["breakdown"]["experience_score"] is None, "Experience score should be None for fresher"
    print(f"[PASS] Fresher Analysis: ATS Score = {score['overall_score']}, Rating = {score['rating']}, Experience Section Hidden = {not exp['include_experience_section']}")


def test_experienced_analysis():
    exp = classify_experience_text(
        ["Senior Software Engineer - Acme Cloud Solutions (2021 - 2024)", "Software Developer - Core Systems Inc (2018 - 2021)"],
        [],
        EXPERIENCED_RESUME_TEXT
    )
    assert exp["candidate_type"] == "experienced", f"Expected experienced, got {exp['candidate_type']}"
    assert exp["has_professional_experience"] is True
    assert exp["include_experience_section"] is True

    match = compare_resume_with_jd(EXPERIENCED_RESUME_TEXT, ["Python", "FastAPI", "Docker", "AWS", "PostgreSQL"], SAMPLE_JD)
    score = calculate_ats_score(
        name="Sarah Jenkins",
        email="sarah.jenkins@example.com",
        phone="+1 (555) 987-6543",
        skills=["Python", "FastAPI", "Docker", "AWS", "PostgreSQL"],
        education=["B.S. in Computer Science"],
        projects=[],
        experience=["Senior Software Engineer - Acme Cloud Solutions", "Software Developer - Core Systems Inc"],
        certifications=[],
        raw_text=EXPERIENCED_RESUME_TEXT,
        jd_text=SAMPLE_JD,
        match_results=match,
        exp_classification=exp,
    )
    assert 0 <= score["overall_score"] <= 100
    assert score["breakdown"]["experience_score"] is not None
    print(f"[PASS] Experienced Analysis: ATS Score = {score['overall_score']}, Rating = {score['rating']}, Experience Score = {score['breakdown']['experience_score']}")


def test_virtual_simulation_analysis():
    exp = classify_experience_text(
        ["JPMorgan Chase & Co. Software Engineering Virtual Experience on Forage (2023)"],
        ["Task Manager App"],
        VIRTUAL_SIMULATION_RESUME_TEXT
    )
    assert exp["has_virtual_experience"] is True, "Expected virtual experience to be detected"
    assert exp["has_professional_experience"] is False, "Virtual simulation should not be classified as direct full-time professional experience"
    assert exp["candidate_type"] == "fresher"
    print(f"[PASS] Virtual Simulation Analysis: Virtual Exp Detected = {exp['has_virtual_experience']}, Candidate Type = {exp['candidate_type']}")


async def test_endpoint_pipeline():
    # Test valid submission with dummy docx or pdf
    import docx
    doc = docx.Document()
    doc.add_heading("Alex Mercer", level=1)
    doc.add_paragraph("alex.mercer@example.com | 555-234-5678")
    doc.add_heading("Skills", level=2)
    doc.add_paragraph("Python, FastAPI, PostgreSQL, Docker, AWS, Git, React")
    doc.add_heading("Projects", level=2)
    doc.add_paragraph("AI Resume Analyzer - Built using FastAPI and React with 95% accuracy.")
    doc.add_heading("Education", level=2)
    doc.add_paragraph("B.S. in Computer Science (2024)")
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_bytes = doc_io.getvalue()

    upload_file = UploadFile(
        file=io.BytesIO(doc_bytes),
        filename="alex_resume.docx",
        headers=Headers({"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"})
    )

    response = await analyze_resume(resume=upload_file, job_description=SAMPLE_JD)
    assert response.success is True
    assert response.ats_score is not None
    assert response.ats_score.overall_score > 0
    assert response.skill_comparison is not None
    assert "FastAPI" in response.skill_comparison.matching_skills
    assert response.job_description_provided is True
    print(f"[PASS] Endpoint Pipeline (With JD): ATS Score = {response.ats_score.overall_score}")

    # Test Optional JD (Resume analysis without JD)
    upload_file_empty_jd = UploadFile(
        file=io.BytesIO(doc_bytes),
        filename="alex_resume.docx",
        headers=Headers({"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"})
    )
    response_no_jd = await analyze_resume(resume=upload_file_empty_jd, job_description="")
    assert response_no_jd.success is True
    assert response_no_jd.job_description_provided is False
    assert response_no_jd.ats_score is not None
    assert response_no_jd.ats_score.overall_score > 0
    print(f"[PASS] Optional JD Analysis: Successfully analyzed without JD (ATS Score = {response_no_jd.ats_score.overall_score})")

    # Test Missing Resume validation
    from fastapi import HTTPException
    try:
        await analyze_resume(resume=None, job_description="Sample JD")
        assert False, "Should have raised HTTPException for missing resume file"
    except HTTPException as exc:
        assert exc.status_code == 400
        print(f"[PASS] Missing Resume Validation: Rejected with status {exc.status_code}")


if __name__ == "__main__":
    test_fresher_analysis()
    test_experienced_analysis()
    test_virtual_simulation_analysis()
    asyncio.run(test_endpoint_pipeline())
    print("\nALL BACKEND UNIT TESTS PASSED SUCCESSFULLY!")

