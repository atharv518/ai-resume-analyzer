import io
import asyncio
from fastapi import UploadFile
from starlette.datastructures import Headers

from app.routes.analyze import analyze_resume
from app.services.experience_detector import classify_experience_text
from app.services.job_matcher import (
    compare_resume_with_jd,
    extract_skills_from_text,
    extract_structured_jd_info,
    check_term_in_text,
    get_canonical_name,
)
from app.services.ats_scorer import calculate_ats_score
from app.services.ai_analyzer import analyze_with_ai, generate_fallback_analysis


SAMPLE_RESUME_TEXT = """
Alex Mercer
alex.mercer@example.com
(555) 234-5678

EDUCATION
Bachelor of Science in Computer Science, Tech State University (2020 - 2024)

SKILLS
Python, FastAPI, JS, React.js, Postgres, Docker, Git, ML

PROJECTS
AI Resume Analyzer
Built full-stack ATS resume analysis platform with FastAPI, React, and PostgreSQL. Implemented text extraction from PDF documents and automated keyword scoring. Reduced parsing latency by 40% and served 200+ active users.

E-Commerce API Service
Developed high-throughput RESTful API with FastAPI and Redis caching, serving 200+ product endpoints with under 50ms latency.

CERTIFICATIONS
AWS Certified Cloud Practitioner
"""

SAMPLE_JD_WITH_SYNONYMS = """
Job Title: Backend / Full-Stack Engineer
Requirements:
- 2+ years of experience with Python and FastAPI
- Strong knowledge of JavaScript, React, and PostgreSQL
- Hands-on experience with Kubernetes (K8s) or Docker containers
- Understanding of Machine Learning (ML) pipelines and REST APIs
- Bachelor's degree in Computer Science or related field
Responsibilities:
- Design and develop scalable microservices and REST APIs
- Optimize database queries and improve system latency
- Collaborate with frontend engineers to build responsive web applications
"""


def test_synonym_and_alias_mapping():
    """Verify that synonyms and tech aliases are properly recognized without false negatives."""
    # Test canonical resolution
    assert get_canonical_name("react.js") == "React"
    assert get_canonical_name("js") == "JavaScript"
    assert get_canonical_name("postgres") == "PostgreSQL"
    assert get_canonical_name("k8s") == "Kubernetes"
    assert get_canonical_name("ml") == "Machine Learning"
    assert get_canonical_name("nlp") == "Natural Language Processing"
    assert get_canonical_name("aws") == "AWS"
    assert get_canonical_name("gcp") == "GCP"

    # Test detection in text
    has_react, _ = check_term_in_text("React", "experienced with react.js and reactjs")
    assert has_react is True

    has_postgres, _ = check_term_in_text("PostgreSQL", "proficient in postgres and psql databases")
    assert has_postgres is True

    has_k8s, _ = check_term_in_text("Kubernetes", "deployed apps on k8s clusters")
    assert has_k8s is True

    # Test compare_resume_with_jd synonym matching
    match = compare_resume_with_jd(
        resume_text=SAMPLE_RESUME_TEXT,
        resume_skills=["Python", "FastAPI", "JS", "React.js", "Postgres", "Docker", "Git", "ML"],
        jd_text=SAMPLE_JD_WITH_SYNONYMS,
    )

    # Resume has JS -> JD has JavaScript
    assert "JavaScript" in match["matching_skills"] or "JS" in match["matching_skills"]
    # Resume has React.js -> JD has React
    assert "React" in match["matching_skills"] or "React.js" in match["matching_skills"]
    # Resume has Postgres -> JD has PostgreSQL
    assert "PostgreSQL" in match["matching_skills"] or "Postgres" in match["matching_skills"]
    # Resume has ML -> JD has Machine Learning
    assert "Machine Learning" in match["matching_skills"] or "ML" in match["matching_skills"]

    assert match["skill_match_percentage"] > 60.0
    print(f"[PASS] Synonym & Alias Mapping: Match % = {match['skill_match_percentage']}%")


def test_extract_skills_from_text_with_aliases():
    """Verify that extract_skills_from_text accurately extracts canonical skills from raw text."""
    extracted = extract_skills_from_text("Experience with React.js, Python, Postgres, K8s, and ML pipelines.")
    assert "React" in extracted
    assert "Python" in extracted
    assert "PostgreSQL" in extracted
    assert "Kubernetes" in extracted
    assert "Machine Learning" in extracted
    print(f"[PASS] extract_skills_from_text: Extracted {len(extracted)} canonical skills")


def test_structured_jd_info_extraction():
    """Verify JD extraction for experience, education, and responsibilities."""
    jd_info = extract_structured_jd_info(SAMPLE_JD_WITH_SYNONYMS)
    assert jd_info["experience_req"] is not None
    assert "2" in jd_info["experience_req"]
    assert jd_info["education_req"] is not None
    assert "Bachelor" in jd_info["education_req"]
    assert len(jd_info["responsibilities"]) >= 2
    print(f"[PASS] Structured JD Info: Exp = '{jd_info['experience_req']}', Edu = '{jd_info['education_req']}', Responsibilities = {len(jd_info['responsibilities'])}")


def test_ats_scorer_phase4():
    """Verify calculate_ats_score with Phase 4 matching results and breakdown."""
    exp = classify_experience_text([], ["AI Resume Analyzer"], SAMPLE_RESUME_TEXT)
    match = compare_resume_with_jd(
        SAMPLE_RESUME_TEXT,
        ["Python", "FastAPI", "JS", "React.js", "Postgres", "Docker", "Git", "ML"],
        SAMPLE_JD_WITH_SYNONYMS,
    )
    score_res = calculate_ats_score(
        name="Alex Mercer",
        email="alex.mercer@example.com",
        phone="(555) 234-5678",
        skills=["Python", "FastAPI", "JS", "React.js", "Postgres", "Docker", "Git", "ML"],
        education=["Bachelor of Science in Computer Science"],
        projects=["AI Resume Analyzer"],
        experience=[],
        certifications=["AWS Certified Cloud Practitioner"],
        raw_text=SAMPLE_RESUME_TEXT,
        jd_text=SAMPLE_JD_WITH_SYNONYMS,
        match_results=match,
        exp_classification=exp,
    )
    assert 0 <= score_res["overall_score"] <= 100
    assert "skills_score" in score_res["breakdown"]
    assert "keyword_score" in score_res["breakdown"]
    assert "projects_score" in score_res["breakdown"]
    assert "education_score" in score_res["breakdown"]
    assert "structure_score" in score_res["breakdown"]
    print(f"[PASS] calculate_ats_score: Score = {score_res['overall_score']}, Rating = {score_res['rating']}")


def test_ai_analyzer_deterministic_fallback():
    """Verify that the deterministic fallback engine populates all Phase 4 structured fields."""
    exp = classify_experience_text([], ["AI Resume Analyzer", "E-Commerce API Service"], SAMPLE_RESUME_TEXT)
    match = compare_resume_with_jd(
        SAMPLE_RESUME_TEXT,
        ["Python", "FastAPI", "JS", "React.js", "Postgres", "Docker", "Git", "ML"],
        SAMPLE_JD_WITH_SYNONYMS,
    )

    result = generate_fallback_analysis(
        name="Alex Mercer",
        skills=["Python", "FastAPI", "JS", "React.js", "Postgres", "Docker", "Git", "ML"],
        projects=["AI Resume Analyzer", "E-Commerce API Service"],
        education=["Bachelor of Science in Computer Science"],
        certifications=["AWS Certified Cloud Practitioner"],
        raw_text=SAMPLE_RESUME_TEXT,
        jd_text=SAMPLE_JD_WITH_SYNONYMS,
        match_results=match,
        exp_classification=exp,
    )

    assert result["is_ai_powered"] is False
    assert result["provider_used"] == "deterministic"
    assert len(result["resume_strengths"]) >= 1
    assert len(result["resume_weaknesses"]) >= 1
    assert "overview" in result["match_explanation"]
    assert len(result["match_explanation"]["strongest_match_areas"]) >= 1
    assert len(result["project_evaluations"]) >= 1
    assert result["project_evaluations"][0]["relevance_score"] in ["High", "Medium", "Low", "Not Relevant"]
    assert len(result["prioritized_recommendations"]["high_priority"]) >= 1 or len(result["prioritized_recommendations"]["medium_priority"]) >= 1
    assert len(result["ats_optimization_tips"]) >= 1
    print(f"[PASS] Deterministic AI Fallback Engine: Evaluated {len(result['project_evaluations'])} projects with complete Phase 4 schema")


async def test_analyze_with_ai_fallback():
    """Verify analyze_with_ai executes gracefully and returns full schema when offline/no API key."""
    exp = classify_experience_text([], ["AI Resume Analyzer"], SAMPLE_RESUME_TEXT)
    match = compare_resume_with_jd(
        SAMPLE_RESUME_TEXT,
        ["Python", "FastAPI", "JS", "React.js", "Postgres", "Docker", "Git", "ML"],
        SAMPLE_JD_WITH_SYNONYMS,
    )
    res = await analyze_with_ai(
        name="Alex Mercer",
        skills=["Python", "FastAPI", "JS", "React.js", "Postgres", "Docker", "Git", "ML"],
        education=["Bachelor of Science in Computer Science"],
        experience=[],
        projects=["AI Resume Analyzer"],
        certifications=["AWS Certified Cloud Practitioner"],
        raw_text=SAMPLE_RESUME_TEXT,
        jd_text=SAMPLE_JD_WITH_SYNONYMS,
        match_results=match,
        exp_classification=exp,
    )
    assert res is not None
    assert "match_explanation" in res
    assert len(res["project_evaluations"]) >= 1
    print(f"[PASS] analyze_with_ai: Provider = {res.get('provider_used')}, AI-Powered = {res.get('is_ai_powered')}")


async def test_end_to_end_analyze_pipeline():
    """Verify the /api/analyze route with full Phase 4 payload (DOCX and PDF)."""
    import docx
    doc = docx.Document()
    doc.add_heading("Alex Mercer", level=1)
    doc.add_paragraph("alex.mercer@example.com | 555-234-5678")
    doc.add_heading("Skills", level=2)
    doc.add_paragraph("Python, FastAPI, React.js, PostgreSQL, Docker, AWS, Git, Machine Learning")
    doc.add_heading("Projects", level=2)
    doc.add_paragraph("AI Resume Analyzer: Built full-stack ATS system using FastAPI and React, improving processing speed by 35%.")
    doc.add_heading("Education", level=2)
    doc.add_paragraph("B.S. in Computer Science (2024)")
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_bytes = doc_io.getvalue()

    upload_file = UploadFile(
        file=io.BytesIO(doc_bytes),
        filename="alex_resume.docx",
        headers=Headers({"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"})
    )

    response = await analyze_resume(resume=upload_file, job_description=SAMPLE_JD_WITH_SYNONYMS)
    assert response.success is True
    assert response.ats_score is not None
    assert response.ats_score.overall_score > 0
    assert response.ai_insights is not None
    assert response.ai_insights.match_explanation is not None
    assert len(response.ai_insights.project_evaluations) >= 1
    assert response.ai_insights.prioritized_recommendations is not None
    # pyrefly: ignore [missing-attribute]
    assert response.skill_comparison.synonym_matches is not None

    print(f"[PASS] End-to-End Pipeline (With JD): Overall ATS = {response.ats_score.overall_score}, Projects Analyzed = {len(response.ai_insights.project_evaluations)}")

    # Test without JD (General resume audit)
    upload_file_no_jd = UploadFile(
        file=io.BytesIO(doc_bytes),
        filename="alex_resume.docx",
        headers=Headers({"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"})
    )
    response_no_jd = await analyze_resume(resume=upload_file_no_jd, job_description="")
    assert response_no_jd.success is True
    assert response_no_jd.job_description_provided is False
    # pyrefly: ignore [missing-attribute]
    assert response_no_jd.ats_score.overall_score > 0
    assert response_no_jd.ai_insights is not None
    # pyrefly: ignore [missing-attribute]
    print(f"[PASS] End-to-End Pipeline (Without JD): Overall ATS = {response_no_jd.ats_score.overall_score}")


async def main():
    test_synonym_and_alias_mapping()
    test_extract_skills_from_text_with_aliases()
    test_structured_jd_info_extraction()
    test_ats_scorer_phase4()
    test_ai_analyzer_deterministic_fallback()
    await test_analyze_with_ai_fallback()
    await test_end_to_end_analyze_pipeline()
    print("\nALL PHASE 4 BACKEND TESTS PASSED SUCCESSFULLY!")


if __name__ == "__main__":
    asyncio.run(main())
