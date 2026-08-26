import io
import logging
import re
from typing import Any
import docx
import pypdf
from fastapi import HTTPException, status

logger = logging.getLogger(__name__)


def clean_linkedin_pdf_artifacts(text: str) -> str:
    """Clean common artifacts from LinkedIn PDF exports (e.g. 'Page X of Y', contact header banners)."""
    if not text:
        return ""
    # Remove "Page X of Y" pagination lines
    cleaned = re.sub(r"(?i)\bpage\s+\d+\s+of\s+\d+\b", "", text)
    # Remove LinkedIn contact header metadata
    cleaned = re.sub(r"(?i)contact\s+www\.linkedin\.com/in/[^\s]+", "", cleaned)
    # Remove "(LinkedIn)" footer annotations
    cleaned = re.sub(r"(?i)\(LinkedIn\)", "", cleaned)
    return cleaned.strip()


def is_suspicious_extraction(text: str) -> bool:
    """Determine if extracted text is malformed, character-spaced, or fragmented."""
    if not text or not text.strip():
        return True

    clean_str = text.strip()
    tokens = clean_str.split()
    if len(tokens) < 5:
        return True

    lines = [line.strip() for line in clean_str.splitlines() if line.strip()]
    if not lines:
        return True

    # 1. Single character alphabetic tokens ratio
    single_char_tokens = [t for t in tokens if len(t) == 1 and t.isalpha()]
    single_char_ratio = len(single_char_tokens) / len(tokens)

    # 2. Single character lines ratio (vertical character streaming)
    single_char_lines = [l for l in lines if len(l) == 1 and l.isalpha()]
    single_char_line_ratio = len(single_char_lines) / len(lines)

    # 3. Average token length
    avg_token_len = sum(len(t) for t in tokens) / len(tokens)

    # 4. Spaced words pattern: e.g. "A L E X" or "S U M M A R Y" or "P y t h o n"
    spaced_word_matches = re.findall(r"(?:\b[A-Za-z]\s+){3,}[A-Za-z]\b", clean_str)

    # 5. Meaningful word ratio (words with 4+ characters)
    meaningful_words = [t for t in tokens if len(t) >= 4 and t.isalpha()]
    meaningful_word_ratio = len(meaningful_words) / len(tokens)

    # Trigger suspicious flag on clear fragmentation signals:
    if single_char_ratio > 0.35 and len(tokens) >= 15:
        return True

    if single_char_line_ratio > 0.30 and len(lines) >= 6:
        return True

    if avg_token_len < 2.2 and len(tokens) >= 15:
        return True

    if len(spaced_word_matches) >= 3:
        return True

    if len(spaced_word_matches) >= 1 and single_char_ratio > 0.20:
        return True

    if len(tokens) >= 25 and meaningful_word_ratio < 0.10:
        return True

    return False


def reconstruct_spaced_text(text: str) -> str:
    """Reconstruct words and sentences from character-spaced or fragmented PDF text."""
    if not text:
        return ""

    lines = text.splitlines()
    reconstructed_lines: list[str] = []

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            reconstructed_lines.append("")
            continue

        tokens = line.split()
        if not tokens:
            continue

        # Check if line looks character-spaced
        single_chars = sum(1 for t in tokens if len(t) == 1 and t.isalnum())
        is_line_spaced = len(tokens) >= 2 and (single_chars / len(tokens)) >= 0.35

        if is_line_spaced:
            # Multi-space (>= 2 spaces) or tabs typically separate words in spaced PDF streams
            word_chunks = re.split(r"\s{2,}|\t+", line)
            cleaned_chunks: list[str] = []

            for chunk in word_chunks:
                chunk = chunk.strip()
                if not chunk:
                    continue

                chunk_tokens = chunk.split(" ")
                chunk_single_chars = sum(1 for t in chunk_tokens if len(t) == 1 and t.isalnum())

                if len(chunk_tokens) >= 2 and (chunk_single_chars / len(chunk_tokens)) >= 0.35:
                    # Merge spaced characters into contiguous words
                    merged = re.sub(r"(?<=[A-Za-z0-9])\s(?=[A-Za-z0-9])", "", chunk)
                    # Fix symbols attached to words (e.g. "Pune ," -> "Pune,", "C + +" -> "C++")
                    merged = re.sub(r"\s+([,.:;|\-–—!?)\]}>])", r"\1", merged)
                    merged = re.sub(r"([(\[{<])\s+", r"\1", merged)
                    merged = re.sub(r"\s*([@])\s*", r"\1", merged)
                    merged = re.sub(r"(?<=[A-Za-z0-9])\s*\+\s*\+", "++", merged)
                    merged = re.sub(r"(?<=\w)\s*\.\s*(?=[A-Za-z])", ".", merged)
                    merged = re.sub(r"(?<=[0-9])\s*\.\s*(?=[0-9])", ".", merged)
                    merged = re.sub(r"(?<=[A-Za-z0-9])\s*-\s*(?=[A-Za-z0-9])", "-", merged)
                    cleaned_chunks.append(merged)
                else:
                    cleaned_chunks.append(chunk)

            reconstructed_lines.append(" ".join(cleaned_chunks))
        else:
            # Even if the whole line is not spaced, fix inline spaced symbols like "C + +" or "C #"
            fixed_line = re.sub(r"(?<=[A-Za-z0-9])\s*\+\s*\+", "++", line)
            fixed_line = re.sub(r"(?<=[A-Za-z0-9])\s*#\b", "#", fixed_line)
            reconstructed_lines.append(fixed_line)

    result = "\n".join(reconstructed_lines)

    # General pass for remaining spaced symbols and standard abbreviations
    result = re.sub(r"\bC\s*\+\s*\+", "C++", result)
    result = re.sub(r"\bC\s*#", "C#", result)

    # Handle vertical single-character lines (e.g. A\nL\nE\nX\n\nJ\nO\nH\nN)
    res_lines = result.splitlines()
    single_char_run: list[str] = []
    final_lines: list[str] = []

    for l in res_lines:
        s = l.strip()
        if len(s) == 1 and s.isalnum():
            single_char_run.append(s)
        else:
            if single_char_run:
                if len(single_char_run) >= 3:
                    final_lines.append("".join(single_char_run))
                else:
                    final_lines.extend(single_char_run)
                single_char_run = []
            final_lines.append(l)

    if single_char_run:
        if len(single_char_run) >= 3:
            final_lines.append("".join(single_char_run))
        else:
            final_lines.extend(single_char_run)

    return "\n".join(final_lines)


def render_pdf_pages_to_images(file_bytes: bytes) -> list[Any]:
    """Render each page of a PDF into a high-resolution PIL Image using pypdfium2."""
    images: list[Any] = []
    try:
        import pypdfium2

        pdf = pypdfium2.PdfDocument(file_bytes)
        for page_idx in range(len(pdf)):
            page = pdf[page_idx]
            # Render at 2x scale (~144 DPI) for crisp OCR character recognition
            pil_image = page.render(scale=2.0).to_pil()
            images.append(pil_image)
    except Exception as exc:
        logger.warning(f"Could not render PDF pages via pypdfium2: {exc}")

    return images


def extract_ocr_from_pdf(file_bytes: bytes, reader: pypdf.PdfReader | None = None) -> str:
    """Attempt OCR extraction on rendered PDF pages (or embedded page images) if OCR engine is available."""
    try:
        import pytesseract
        from PIL import Image

        extracted_text_pages: list[str] = []

        # Strategy A: Render entire PDF pages as images (recovers vector layout, character-spaced fonts, and flattened scans)
        rendered_images = render_pdf_pages_to_images(file_bytes)
        for img in rendered_images:
            try:
                ocr_result = pytesseract.image_to_string(img)
                ocr_text = _normalize_ocr_output(ocr_result)
                if ocr_text.strip():
                    extracted_text_pages.append(ocr_text.strip())
            except Exception as ocr_err:
                logger.warning(f"OCR page extraction failed: {ocr_err}")
                continue

        # Strategy B: Fallback to embedded image extraction if page rendering yielded nothing
        if not extracted_text_pages and reader is not None:
            for page in reader.pages:
                for image_file in getattr(page, "images", []):
                    data = getattr(image_file, "data", None)
                    if data is None and isinstance(image_file, dict):
                        data = image_file.get("data")
                    if not data:
                        continue

                    img = Image.open(io.BytesIO(data))
                    ocr_result = pytesseract.image_to_string(img)
                    ocr_text = _normalize_ocr_output(ocr_result)
                    if ocr_text.strip():
                        extracted_text_pages.append(ocr_text.strip())

        return "\n\n".join(extracted_text_pages).strip()
    except Exception as exc:
        logger.warning(f"OCR extraction engine unavailable or failed: {exc}")
        return ""


def _normalize_ocr_output(ocr_result: Any) -> str:
    """Normalize varied return types from pytesseract / mocks."""
    if isinstance(ocr_result, dict):
        if "text" in ocr_result:
            val = ocr_result["text"]
            return " ".join(val) if isinstance(val, list) else str(val)
        return " ".join(str(v) for v in ocr_result.values() if v)
    if isinstance(ocr_result, list):
        return " ".join(str(item) for item in ocr_result if item)
    if isinstance(ocr_result, str):
        return ocr_result
    return str(ocr_result) if ocr_result is not None else ""


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract readable text from PDF bytes with quality verification, rendered page OCR, and de-spacing fallbacks."""
    try:
        reader = pypdf.PdfReader(io.BytesIO(file_bytes), strict=False)
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Could not read the PDF file. It may be corrupted or in an unsupported format.",
        ) from exc

    if reader.is_encrypted:
        try:
            # Attempt empty password decrypt if standard permissions
            decrypted = reader.decrypt("")
            if decrypted == 0:
                reader.decrypt(b"")
        except Exception:
            pass  # Some PDF readers still allow reading unencrypted streams

    # Step 1: Standard pypdf extraction
    pages_text: list[str] = []
    try:
        for page in reader.pages:
            try:
                page_text = page.extract_text() or ""
                # If standard extraction is very sparse, attempt layout extraction
                if len(page_text.strip().split()) < 5:
                    try:
                        layout_text = page.extract_text(extraction_mode="layout") or ""
                        if len(layout_text.strip().split()) > len(page_text.strip().split()):
                            page_text = layout_text
                    except Exception:
                        pass
                if page_text.strip():
                    pages_text.append(page_text.strip())
            except Exception:
                continue
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Failed to extract text from the PDF file.",
        ) from exc

    full_text = "\n\n".join(pages_text).strip()
    full_text = full_text.replace("\x00", "")

    # Step 2: Quality Check
    # If pypdf extracted clean, good-quality text, use it directly (fast path)
    if full_text and not is_suspicious_extraction(full_text):
        return clean_linkedin_pdf_artifacts(full_text)

    # Step 3: Fallback Strategy for Suspicious / Spaced / Scanned PDFs
    # Fallback 3A: Attempt rendered-page OCR if OCR engine is available
    ocr_result = extract_ocr_from_pdf(file_bytes, reader)
    if ocr_result and not is_suspicious_extraction(ocr_result):
        return clean_linkedin_pdf_artifacts(ocr_result)

    # Fallback 3B: Attempt intelligent text de-spacing and token reconstruction
    if full_text:
        reconstructed = reconstruct_spaced_text(full_text)
        if reconstructed and not is_suspicious_extraction(reconstructed):
            return clean_linkedin_pdf_artifacts(reconstructed)

    # Fallback 3C: If OCR returned partial text that can also be de-spaced
    if ocr_result:
        reconstructed_ocr = reconstruct_spaced_text(ocr_result)
        if reconstructed_ocr and not is_suspicious_extraction(reconstructed_ocr):
            return clean_linkedin_pdf_artifacts(reconstructed_ocr)

    # Step 4: If all extraction & fallback strategies failed to produce readable text
    raise HTTPException(
        status_code=status.HTTP_400_BAD_REQUEST,
        detail="No readable text could be extracted. The PDF appears to be a scanned image or flattened graphic without a selectable text layer. Please upload a standard text PDF, DOCX, TXT, or RTF file.",
    )


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract readable text from DOCX bytes."""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Could not read the DOCX file. It may be corrupted or invalid.",
        ) from exc

    lines: list[str] = []
    try:
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                lines.append(text)

        # Extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    lines.append(" | ".join(row_text))
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Failed to extract text from the DOCX file.",
        ) from exc

    full_text = "\n".join(lines).strip()
    if not full_text:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No readable text could be extracted from the DOCX file.",
        )

    return full_text


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Extract readable text from plain text (.txt) bytes supporting multiple encodings."""
    encodings = ["utf-8-sig", "utf-8", "utf-16", "latin-1", "cp1252"]
    for enc in encodings:
        try:
            text = file_bytes.decode(enc)
            cleaned = text.strip()
            if cleaned:
                return cleaned
        except (UnicodeDecodeError, ValueError):
            continue

    raise HTTPException(
        status_code=status.HTTP_400_BAD_REQUEST,
        detail="Could not decode the text file. Please ensure it is saved in UTF-8 or standard ASCII format.",
    )


def extract_text_from_rtf(file_bytes: bytes) -> str:
    """Extract clean readable text from Rich Text Format (.rtf) bytes."""
    try:
        raw_str = file_bytes.decode("latin-1", errors="replace")
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Failed to decode RTF document.",
        ) from exc

    # Remove non-text metadata groups
    text = re.sub(r"\\(?:pict|bin|fonttbl|colortbl|stylesheet|info)\b.*?[;}]", "", raw_str, flags=re.DOTALL)

    # Convert Unicode escapes \uN?
    def _replace_unicode(match: re.Match) -> str:
        codepoint = int(match.group(1))
        if codepoint < 0:
            codepoint += 65536
        return chr(codepoint)

    text = re.sub(r"\\u(-?\d+)\??", _replace_unicode, text)

    # Convert hex escapes \'XX
    def _replace_hex(match: re.Match) -> str:
        return bytes.fromhex(match.group(1)).decode("latin-1", errors="replace")

    text = re.sub(r"\\\'([0-9a-fA-F]{2})", _replace_hex, text)

    # Convert line/paragraph breaks
    text = re.sub(r"\\(?:par|line|page)\b", "\n", text)
    text = re.sub(r"\\(?:tab)\b", "\t", text)

    # Remove remaining control words
    text = re.sub(r"\\[a-zA-Z]+-?\d*\s?", "", text)

    # Remove group braces
    text = re.sub(r"[{}]", "", text)

    cleaned_lines = [line.strip() for line in text.splitlines() if line.strip()]
    full_text = "\n".join(cleaned_lines)

    if not full_text:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No readable text could be extracted from the RTF file.",
        )

    return full_text


def extract_resume_text(extension: str, file_bytes: bytes) -> str:
    """Extract text based on file extension (.pdf, .docx, .txt, .rtf)."""
    normalized_ext = extension.lower()
    if normalized_ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    if normalized_ext == ".docx":
        return extract_text_from_docx(file_bytes)
    if normalized_ext == ".txt":
        return extract_text_from_txt(file_bytes)
    if normalized_ext == ".rtf":
        return extract_text_from_rtf(file_bytes)

    raise HTTPException(
        status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
        detail="Only PDF, DOCX, TXT, and RTF resume files are supported.",
    )
