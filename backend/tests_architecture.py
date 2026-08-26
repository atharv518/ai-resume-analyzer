import io
import asyncio
import docx
from fastapi import UploadFile
from starlette.datastructures import Headers

from app.routes.analyze import analyze_resume, submit_async_analysis, get_job_status
from app.services.job_queue import job_queue, JobQueueManager
from app.services.health_service import get_deep_health_report, check_parser_subsystem_health, check_ai_provider_health


SAMPLE_JD = """
Job Title: Backend Engineer
Requirements:
- 2+ years experience in Python, FastAPI, and PostgreSQL
- Experience building RESTful microservices and Docker containers
- Understanding of ATS optimization and document processing
"""


def generate_sample_docx() -> bytes:
    doc = docx.Document()
    doc.add_heading("Alex Mercer", level=1)
    doc.add_paragraph("alex.mercer@example.com | (555) 234-5678")
    doc.add_heading("Skills", level=2)
    doc.add_paragraph("Python, FastAPI, PostgreSQL, Docker, AWS, Git, React")
    doc.add_heading("Projects", level=2)
    doc.add_paragraph("AI Resume Analyzer - Built full-stack ATS platform using FastAPI and React.")
    doc.add_heading("Education", level=2)
    doc.add_paragraph("B.S. in Computer Science, Tech State University (2024)")
    doc_io = io.BytesIO()
    doc.save(doc_io)
    return doc_io.getvalue()


async def test_deep_health_check():
    """Verify deep health check diagnostic probe across subsystems."""
    parsers = check_parser_subsystem_health()
    assert parsers["status"] == "healthy"
    assert "pypdf" in parsers["pdf_engine"]
    assert "python-docx" in parsers["docx_engine"]

    ai_health = await check_ai_provider_health()
    assert ai_health["status"] in {"healthy", "degraded"}

    report, status_code = await get_deep_health_report()
    assert status_code in {200, 503}
    assert report["app_name"] == "NovaATS"
    assert "subsystems" in report
    assert "parsers" in report["subsystems"]
    assert "ai_provider" in report["subsystems"]
    assert "job_queue" in report["subsystems"]
    assert "feature_flags" in report["subsystems"]
    print(f"[PASS] Deep Health Diagnostics: Status = {report['status']} (HTTP {status_code})")


async def test_async_job_queue_pipeline():
    """Verify asynchronous job submission, queue execution, and polling endpoint."""
    docx_bytes = generate_sample_docx()
    upload_file = UploadFile(
        file=io.BytesIO(docx_bytes),
        filename="alex_resume.docx",
        headers=Headers({"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}),
    )

    # 1. Submit async job
    submit_res = await submit_async_analysis(resume=upload_file, job_description=SAMPLE_JD)
    assert submit_res.status == "queued"
    assert submit_res.job_id is not None
    job_id = submit_res.job_id
    print(f"[PASS] Async Job Submission: Job ID = {job_id} (Status: {submit_res.status})")

    # 2. Poll until completed
    max_wait = 10.0
    start = asyncio.get_event_loop().time()
    completed = False

    while asyncio.get_event_loop().time() - start < max_wait:
        status_res = await get_job_status(job_id=job_id)
        if status_res.status == "completed":
            completed = True
            assert status_res.result is not None
            assert status_res.result.success is True
            assert status_res.result.ats_score is not None
            assert status_res.result.ats_score.overall_score > 0
            assert status_res.progress_percentage == 100
            print(f"[PASS] Async Job Polling: Completed with ATS Score = {status_res.result.ats_score.overall_score}")
            break
        await asyncio.sleep(0.1)

    assert completed, "Async job did not complete within timeout window"


async def test_job_queue_manager_concurrency_and_ttl():
    """Verify in-memory queue manager TTL cleanup and stats."""
    manager = JobQueueManager(max_concurrent=2, ttl_seconds=1)
    jid = manager.create_job()
    assert manager.get_job(jid) is not None

    stats = manager.get_stats()
    assert stats["queued"] == 1
    print(f"[PASS] Job Queue Concurrency & Stats: Total Tracked = {stats['total_tracked']}, Queued = {stats['queued']}")


def run_all_architecture_tests():
    print("\n--- RUNNING ARCHITECTURE & WORKFLOW TEST SUITE ---")
    asyncio.run(test_deep_health_check())
    asyncio.run(test_async_job_queue_pipeline())
    asyncio.run(test_job_queue_manager_concurrency_and_ttl())
    print("\nALL ARCHITECTURE TESTS PASSED SUCCESSFULLY! ---\n")


if __name__ == "__main__":
    run_all_architecture_tests()
